#!/usr/bin/env python3
"""
YueShen Knowledge Base RAG MCP Server

- Load pdf/docx from a directory, chunk, and write into Chroma.
- Embedding config is provided via MCP tool params or env vars, not system.conf.
- Auto-ingest watcher can run on startup to keep the index fresh.
- Tools: ingest_yueshen, query_yueshen, yueshen_stats.
"""

import hashlib
import json
import logging
import os
import re
import sys
import time
import threading
from dataclasses import dataclass
from typing import Any, Dict, Iterable, List, Optional, Tuple

import requests

# Keep stdout clean for MCP stdio; route logs to stderr and disable Chroma telemetry noise.
os.environ.setdefault("CHROMA_TELEMETRY", "FALSE")

# Make project root importable (for optional fallback embedding)
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

try:
    from mcp.server import Server
    from mcp.types import Tool, TextContent
    import mcp.server.stdio
except ImportError:
    print("MCP library not installed. Please run: pip install mcp", file=sys.stderr, flush=True)
    sys.exit(1)

try:
    import chromadb
except ImportError:
    print("chromadb not installed. Please run: pip install chromadb", file=sys.stderr, flush=True)
    sys.exit(1)

server = Server("yueshen_rag")

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s - %(message)s",
    stream=sys.stderr,
)
logger = logging.getLogger("yueshen_rag")

# Defaults (can be overridden via env)
DEFAULT_CORPUS_DIR = os.getenv(
    "YUESHEN_CORPUS_DIR",
    os.path.join(PROJECT_ROOT, "新知识库"),
)
DEFAULT_PERSIST_DIR = os.getenv(
    "YUESHEN_PERSIST_DIR",
    os.path.join(PROJECT_ROOT, "cache_data", "chromadb_yueshen"),
)
COLLECTION_NAME = "yueshen_kb"
DEFAULT_EMBED_BASE_URL = os.getenv("YUESHEN_EMBED_BASE_URL")
DEFAULT_EMBED_API_KEY = os.getenv("YUESHEN_EMBED_API_KEY")
DEFAULT_EMBED_MODEL = os.getenv("YUESHEN_EMBED_MODEL", "text-embedding-3-small")
AUTO_INGEST_ENABLED = os.getenv("YUESHEN_AUTO_INGEST", "1") != "0"
AUTO_INGEST_INTERVAL = int(os.getenv("YUESHEN_AUTO_INTERVAL", "300"))
AUTO_RESET_ON_START = os.getenv("YUESHEN_AUTO_RESET_ON_START", "0") != "0"


# -------------------- Text chunking -------------------- #
def _len_with_newlines(parts: List[str]) -> int:
    if not parts:
        return 0
    return sum(len(p) for p in parts) + (len(parts) - 1)


def split_into_chunks(text: str, chunk_size: int = 600, overlap: int = 120) -> List[str]:
    """Paragraph/ sentence-aware chunking with small overlap."""
    cleaned = re.sub(r"[ \t]+", " ", text.replace("\u00a0", " ")).strip()
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", cleaned) if p.strip()]
    segments: List[str] = []
    for para in paragraphs:
        if len(para) <= chunk_size:
            segments.append(para)
        else:
            for sent in re.split(r"(?<=[。！？!?…])", para):
                s = sent.strip()
                if s:
                    segments.append(s)

    chunks: List[str] = []
    buf: List[str] = []
    for seg in segments:
        seg = seg.strip()
        if not seg:
            continue
        current_len = _len_with_newlines(buf)
        if current_len + len(seg) + (1 if buf else 0) <= chunk_size:
            buf.append(seg)
            continue

        if buf:
            chunks.append("\n".join(buf).strip())

        # build overlap from previous chunk tail
        buf = []
        if overlap > 0 and chunks:
            tail: List[str] = []
            tail_len = 0
            for s in reversed(chunks[-1].split("\n")):
                tail.insert(0, s)
                tail_len += len(s)
                if tail_len >= overlap:
                    break
            if tail:
                buf.extend(tail)

        buf.append(seg)

    if buf:
        chunks.append("\n".join(buf).strip())
    return chunks


# -------------------- Document readers -------------------- #
def _extract_docx(path: str) -> str:
    from docx import Document

    doc = Document(path)
    texts: List[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            texts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                texts.append(" | ".join(cells))

    return "\n".join(texts)


def _extract_pdf_pages(path: str) -> List[Tuple[int, str]]:
    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError("pdfplumber is required for pdf parsing") from exc

    pages: List[Tuple[int, str]] = []
    with pdfplumber.open(path) as pdf:
        for idx, page in enumerate(pdf.pages, start=1):
            txt = page.extract_text() or ""
            pages.append((idx, txt))
    return pages


# -------------------- Data models -------------------- #
@dataclass
class Chunk:
    text: str
    source_path: str
    page: Optional[int]
    chunk_id: str
    metadata: Dict[str, Any]


# -------------------- Corpus loader -------------------- #
class CorpusLoader:
    def __init__(self, root_dir: str = DEFAULT_CORPUS_DIR, chunk_size: int = 600, overlap: int = 120):
        self.root_dir = root_dir
        self.chunk_size = chunk_size
        self.overlap = overlap

    def _iter_files(self) -> Iterable[str]:
        for root, _, files in os.walk(self.root_dir):
            for fn in files:
                if fn.lower().endswith((".pdf", ".docx")):
                    yield os.path.join(root, fn)

    def _file_to_chunks(self, path: str) -> List[Chunk]:
        ext = os.path.splitext(path)[1].lower()
        rel_path = os.path.relpath(path, self.root_dir)
        chunks: List[Chunk] = []

        try:
            if ext == ".pdf":
                pages = _extract_pdf_pages(path)
                for page_num, page_text in pages:
                    for idx, chunk_text in enumerate(
                        split_into_chunks(page_text, chunk_size=self.chunk_size, overlap=self.overlap)
                    ):
                        chunk_id = hashlib.md5(
                            f"{rel_path}|{page_num}|{idx}|{chunk_text}".encode("utf-8", errors="ignore")
                        ).hexdigest()
                        chunks.append(
                            Chunk(
                                text=chunk_text,
                                source_path=rel_path,
                                page=page_num,
                                chunk_id=chunk_id,
                                metadata={"source": rel_path, "page": page_num, "ext": ext},
                            )
                        )
            elif ext == ".docx":
                text = _extract_docx(path)
                for idx, chunk_text in enumerate(
                    split_into_chunks(text, chunk_size=self.chunk_size, overlap=self.overlap)
                ):
                    chunk_id = hashlib.md5(
                        f"{rel_path}|docx|{idx}|{chunk_text}".encode("utf-8", errors="ignore")
                    ).hexdigest()
                    chunks.append(
                        Chunk(
                            text=chunk_text,
                            source_path=rel_path,
                            page=None,
                            chunk_id=chunk_id,
                            metadata={"source": rel_path, "ext": ext},
                        )
                    )
        except Exception as exc:
            logger.warning("Skip file due to parse error %s: %s", rel_path, exc)

        return chunks

    def load(self, max_files: Optional[int] = None) -> List[Chunk]:
        all_chunks: List[Chunk] = []
        for idx, file_path in enumerate(self._iter_files(), start=1):
            if max_files and idx > max_files:
                break
            all_chunks.extend(self._file_to_chunks(file_path))
        return all_chunks


# -------------------- Embedding backend -------------------- #
class EmbeddingBackend:
    """Embedding client with API config, falling back to project encoder if needed."""

    def __init__(
        self,
        base_url: Optional[str] = None,
        api_key: Optional[str] = None,
        model: Optional[str] = None,
    ):
        self.base_url = base_url or DEFAULT_EMBED_BASE_URL
        self.api_key = api_key or DEFAULT_EMBED_API_KEY
        self.model = model or DEFAULT_EMBED_MODEL
        self._cache: Dict[str, List[float]] = {}
        self._fallback_encoder = None
        try:
            from simulation_engine.gpt_structure import get_text_embedding as _fallback

            self._fallback_encoder = _fallback
        except Exception as exc:
            logger.info("Fallback embedding not available: %s", exc)

    def _call_api(self, text: str) -> List[float]:
        if not self.base_url or not self.api_key:
            raise RuntimeError("Embedding API config missing (base_url/api_key)")
        url = self.base_url.rstrip("/") + "/embeddings"
        payload = {"input": text, "model": self.model}
        headers = {"Authorization": f"Bearer {self.api_key}"}
        resp = requests.post(url, json=payload, headers=headers, timeout=30)
        if resp.status_code != 200:
            raise RuntimeError(f"Embedding API error: {resp.status_code} {resp.text}")
        data = resp.json()
        embedding = data.get("data", [{}])[0].get("embedding")
        if embedding is None:
            raise RuntimeError("Embedding API response missing embedding")
        return embedding

    def encode(self, text: str) -> List[float]:
        cache_key = hashlib.md5(f"{self.model}|{self.base_url}|{text}".encode("utf-8", errors="ignore")).hexdigest()
        if cache_key in self._cache:
            return self._cache[cache_key]

        embedding: Optional[List[float]] = None
        if self.base_url and self.api_key:
            embedding = self._call_api(text)
        elif self._fallback_encoder:
            embedding = self._fallback_encoder(text)
        else:
            raise RuntimeError("No embedding method available (provide base_url/api_key or enable fallback).")

        if not isinstance(embedding, list):
            embedding = list(embedding)
        self._cache[cache_key] = embedding
        return embedding


# -------------------- Chroma store -------------------- #
class ChromaStore:
    def __init__(
        self,
        persist_dir: str = DEFAULT_PERSIST_DIR,
        collection_name: str = COLLECTION_NAME,
        embedder: Optional[EmbeddingBackend] = None,
    ):
        os.makedirs(persist_dir, exist_ok=True)
        self.persist_dir = persist_dir
        self.collection_name = collection_name
        self.embedder = embedder or EmbeddingBackend()
        self.client = chromadb.PersistentClient(path=persist_dir)
        self.collection = self.client.get_or_create_collection(collection_name)

    def reset(self):
        self.client.delete_collection(self.collection_name)
        self.collection = self.client.get_or_create_collection(self.collection_name)

    def upsert_chunks(self, chunks: List[Chunk], batch_size: int = 32) -> Dict[str, Any]:
        start = time.time()
        total = 0
        ids: List[str] = []
        docs: List[str] = []
        metas: List[Dict[str, Any]] = []
        embs: List[List[float]] = []

        def flush():
            nonlocal total, ids, docs, metas, embs
            if not ids:
                return
            self.collection.upsert(ids=ids, documents=docs, metadatas=metas, embeddings=embs)
            total += len(ids)
            ids, docs, metas, embs = [], [], [], []

        for chunk in chunks:
            ids.append(chunk.chunk_id)
            docs.append(chunk.text)
            metas.append(chunk.metadata)
            try:
                embs.append(self.embedder.encode(chunk.text))
            except Exception as exc:
                logger.error("Embedding failed, skip id=%s: %s", chunk.chunk_id, exc)
                ids.pop()
                docs.pop()
                metas.pop()
                continue

            if len(ids) >= batch_size:
                flush()

        flush()
        elapsed = time.time() - start
        return {"inserted": total, "seconds": round(elapsed, 2)}

    def query(self, query: str, top_k: int = 5, where: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        emb = self.embedder.encode(query)
        res = self.collection.query(query_embeddings=[emb], n_results=top_k, where=where if where else None)
        results = []
        ids = res.get("ids", [[]])[0]
        docs = res.get("documents", [[]])[0]
        metas = res.get("metadatas", [[]])[0]
        dists = res.get("distances", [[]])[0] if "distances" in res else [None] * len(ids)
        for i in range(len(ids)):
            results.append(
                {"id": ids[i], "document": docs[i], "metadata": metas[i], "distance": dists[i]}
            )
        return {"results": results, "count": len(results)}

    def stats(self) -> Dict[str, Any]:
        try:
            count = self.collection.count()
        except Exception:
            count = None
        return {"persist_dir": self.persist_dir, "collection": self.collection_name, "vectors": count}


# -------------------- Knowledge manager -------------------- #
class KnowledgeManager:
    def __init__(
        self,
        corpus_dir: str = DEFAULT_CORPUS_DIR,
        persist_dir: str = DEFAULT_PERSIST_DIR,
        embedder: Optional[EmbeddingBackend] = None,
    ):
        self.corpus_dir = corpus_dir
        self.persist_dir = persist_dir
        self.embedder = embedder or EmbeddingBackend()
        self.store = ChromaStore(persist_dir=persist_dir, collection_name=COLLECTION_NAME, embedder=self.embedder)

    def _refresh_embedder(self, base_url: Optional[str], api_key: Optional[str], model: Optional[str]):
        if any([base_url, api_key, model]):
            self.embedder = EmbeddingBackend(base_url=base_url, api_key=api_key, model=model)
            self.store = ChromaStore(
                persist_dir=self.persist_dir, collection_name=COLLECTION_NAME, embedder=self.embedder
            )

    def ingest(
        self,
        corpus_dir: Optional[str] = None,
        reset: bool = False,
        chunk_size: int = 600,
        overlap: int = 120,
        batch_size: int = 32,
        max_files: Optional[int] = None,
        embedding_base_url: Optional[str] = None,
        embedding_api_key: Optional[str] = None,
        embedding_model: Optional[str] = None,
    ) -> Dict[str, Any]:
        self._refresh_embedder(embedding_base_url, embedding_api_key, embedding_model)
        target_dir = corpus_dir or self.corpus_dir
        loader = CorpusLoader(root_dir=target_dir, chunk_size=chunk_size, overlap=overlap)
        if reset:
            self.store.reset()

        chunks = loader.load(max_files=max_files)
        logger.info("Loaded %d chunks from %s, start upsert...", len(chunks), target_dir)
        upsert_res = self.store.upsert_chunks(chunks, batch_size=batch_size)
        return {
            "success": True,
            "message": "ingest completed",
            "chunks": len(chunks),
            "inserted": upsert_res.get("inserted", 0),
            "seconds": upsert_res.get("seconds"),
            "persist_dir": self.persist_dir,
            "collection": COLLECTION_NAME,
            "corpus_dir": target_dir,
            "embedding_base_url": self.embedder.base_url,
            "embedding_model": self.embedder.model,
        }

    def query(
        self,
        query: str,
        top_k: int = 5,
        where: Optional[Dict[str, Any]] = None,
        embedding_base_url: Optional[str] = None,
        embedding_api_key: Optional[str] = None,
        embedding_model: Optional[str] = None,
    ) -> Dict[str, Any]:
        self._refresh_embedder(embedding_base_url, embedding_api_key, embedding_model)
        return self.store.query(query=query, top_k=top_k, where=where)

    def stats(self) -> Dict[str, Any]:
        info = self.store.stats()
        info.update(
            {
                "default_corpus_dir": self.corpus_dir,
                "embedding_base_url": self.embedder.base_url,
                "embedding_model": self.embedder.model,
            }
        )
        return info


manager = KnowledgeManager()


# -------------------- Auto ingest watcher -------------------- #
class AutoIngestor:
    """Simple polling-based watcher to auto-ingest when files change."""

    def __init__(
        self,
        km: KnowledgeManager,
        interval_sec: int = AUTO_INGEST_INTERVAL,
        reset_on_start: bool = AUTO_RESET_ON_START,
        enabled: bool = AUTO_INGEST_ENABLED,
    ):
        self.km = km
        self.interval = max(30, interval_sec)
        self.reset_on_start = reset_on_start
        self.enabled = enabled
        self._stop = threading.Event()
        self._thread: Optional[threading.Thread] = None
        self._snapshot: Dict[str, Tuple[float, int]] = {}

    def _take_snapshot(self) -> Dict[str, Tuple[float, int]]:
        snap: Dict[str, Tuple[float, int]] = {}
        for root, _, files in os.walk(self.km.corpus_dir):
            for fn in files:
                if fn.lower().endswith((".pdf", ".docx")):
                    path = os.path.join(root, fn)
                    try:
                        st = os.stat(path)
                        snap[path] = (st.st_mtime, st.st_size)
                    except OSError:
                        continue
        return snap

    def _has_changes(self) -> bool:
        new_snap = self._take_snapshot()
        if new_snap != self._snapshot:
            self._snapshot = new_snap
            return True
        return False

    def _ingest_once(self, reset: bool = False):
        try:
            res = self.km.ingest(
                corpus_dir=self.km.corpus_dir,
                reset=reset,
                embedding_base_url=self.km.embedder.base_url,
                embedding_api_key=self.km.embedder.api_key,
                embedding_model=self.km.embedder.model,
            )
            logger.info("Auto-ingest done: %s", json.dumps(res, ensure_ascii=False))
        except Exception as exc:
            logger.error("Auto-ingest failed: %s", exc)

    def _loop(self):
        # initial snapshot and optional first ingest
        self._snapshot = self._take_snapshot()
        if self.reset_on_start:
            logger.info("Auto-ingest on start (reset=%s)...", self.reset_on_start)
            self._ingest_once(reset=True)
        elif self.enabled:
            logger.info("Auto-ingest initial run...")
            self._ingest_once(reset=False)

        while not self._stop.wait(self.interval):
            if self._has_changes():
                logger.info("Detected corpus change, auto-ingest...")
                self._ingest_once(reset=False)

    def start(self):
        if not self.enabled:
            logger.info("Auto-ingest disabled via env (YUESHEN_AUTO_INGEST=0)")
            return
        if self._thread and self._thread.is_alive():
            return
        self._thread = threading.Thread(target=self._loop, daemon=True)
        self._thread.start()
        logger.info("Auto-ingest watcher started, interval=%ss", self.interval)

    def stop(self):
        self._stop.set()
        if self._thread:
            self._thread.join(timeout=2)


# -------------------- Skip patterns for trivial queries -------------------- #
SKIP_QUERY_PATTERNS = [
    # 问候语
    r'^你好[啊呀吗]?$', r'^hello[!！]?$', r'^hi[!！]?$', r'^嗨[!！]?$', r'^hey[!！]?$',
    r'^早[上]?好[啊呀]?$', r'^晚[上]?好[啊呀]?$', r'^下午好[啊呀]?$', r'^中午好[啊呀]?$',
    # 简单回复
    r'^ok[!！]?$', r'^好[的吧啊呀]?[!！]?$', r'^行[!！]?$', r'^可以[!！]?$', r'^没问题[!！]?$',
    r'^嗯[嗯]?[!！]?$', r'^哦[哦]?[!！]?$', r'^噢[!！]?$',
    # 笑声/情绪
    r'^哈哈[哈]*[!！]?$', r'^呵呵[呵]*[!！]?$', r'^嘿嘿[嘿]*[!！]?$', r'^嘻嘻[嘻]*[!！]?$',
    r'^哼[!！]?$', r'^呜呜[呜]*[!！]?$',
    # 日常用语
    r'^睡觉[了去]?[!！]?$', r'^晚安[!！]?$', r'^再见[!！]?$', r'^拜拜[!！]?$', r'^bye[!！]?$',
    r'^谢谢[你您]?[!！]?$', r'^感谢[!！]?$', r'^thanks[!！]?$', r'^thank you[!！]?$',
    r'^对不起[!！]?$', r'^抱歉[!！]?$', r'^sorry[!！]?$',
    r'^是[的吧啊]?[!！]?$', r'^对[的吧啊]?[!！]?$', r'^不是[!！]?$', r'^不对[!！]?$',
    r'^知道了[!！]?$', r'^明白了[!！]?$', r'^懂了[!！]?$', r'^了解[!！]?$',
    r'^收到[!！]?$', r'^好嘞[!！]?$', r'^得嘞[!！]?$',
    # 疑问简单回复
    r'^啥[?？]?$', r'^什么[?？]?$', r'^嗯[?？]$', r'^哈[?？]$',
    # 单字或极短
    r'^[.。,，!！?？~～]+$',
]

def _is_trivial_query(query: str) -> bool:
    """Check if query is a trivial greeting or simple response that doesn't need KB search."""
    if not query:
        return True
    q = query.strip().lower()
    if len(q) <= 2:
        return True
    for pattern in SKIP_QUERY_PATTERNS:
        if re.match(pattern, q, re.IGNORECASE):
            return True
    return False


# -------------------- MCP tools -------------------- #
@server.list_tools()
async def list_tools() -> List[Tool]:
    return [
        Tool(
            name="ingest_yueshen",
            description="Scan directory (pdf/docx/doc), chunk and write to Chroma",
            inputSchema={
                "type": "object",
                "properties": {
                    "corpus_dir": {"type": "string", "description": "Optional corpus directory override"},
                    "reset": {"type": "boolean", "description": "Recreate collection before ingest", "default": False},
                    "chunk_size": {"type": "integer", "description": "Chunk length (chars)", "default": 600},
                    "overlap": {"type": "integer", "description": "Chunk overlap (chars)", "default": 120},
                    "batch_size": {"type": "integer", "description": "Upsert batch size", "default": 32},
                    "max_files": {"type": "integer", "description": "Optional limit for quick test"},
                    "embedding_base_url": {"type": "string", "description": "Embedding API base url"},
                    "embedding_api_key": {"type": "string", "description": "Embedding API key"},
                    "embedding_model": {"type": "string", "description": "Embedding model name"},
                },
            },
        ),
        Tool(
            name="query_yueshen",
            description="Vector search in YueShen KB",
            inputSchema={
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "User query"},
                    "top_k": {"type": "integer", "description": "Number of results", "default": 5},
                    "where": {"type": "object", "description": "Optional metadata filter (Chroma where)"},
                    "embedding_base_url": {"type": "string", "description": "Embedding API base url"},
                    "embedding_api_key": {"type": "string", "description": "Embedding API key"},
                    "embedding_model": {"type": "string", "description": "Embedding model name"},
                },
                "required": ["query"],
            },
        ),
        Tool(
            name="yueshen_stats",
            description="Show current vector store stats",
            inputSchema={"type": "object", "properties": {}},
        ),
    ]


@server.call_tool()
async def call_tool(name: str, arguments: Dict[str, Any]) -> List[TextContent]:
    try:
        if name == "ingest_yueshen":
            res = manager.ingest(
                corpus_dir=arguments.get("corpus_dir"),
                reset=bool(arguments.get("reset", False)),
                chunk_size=int(arguments.get("chunk_size", 600)),
                overlap=int(arguments.get("overlap", 120)),
                batch_size=int(arguments.get("batch_size", 32)),
                max_files=arguments.get("max_files"),
                embedding_base_url=arguments.get("embedding_base_url"),
                embedding_api_key=arguments.get("embedding_api_key"),
                embedding_model=arguments.get("embedding_model"),
            )
            return [TextContent(type="text", text=json.dumps(res, ensure_ascii=False, indent=2))]

        if name == "query_yueshen":
            query_text = arguments.get("query", "")
            # 跳过常见问候和简单回复，不进行知识库查询
            if _is_trivial_query(query_text):
                return [TextContent(type="text", text=json.dumps({
                    "results": [],
                    "count": 0,
                    "skipped": True,
                    "reason": "trivial query (greeting or simple response)"
                }, ensure_ascii=False, indent=2))]
            res = manager.query(
                query=query_text,
                top_k=int(arguments.get("top_k", 5)),
                where=arguments.get("where"),
                embedding_base_url=arguments.get("embedding_base_url"),
                embedding_api_key=arguments.get("embedding_api_key"),
                embedding_model=arguments.get("embedding_model"),
            )
            return [TextContent(type="text", text=json.dumps(res, ensure_ascii=False, indent=2))]

        if name == "yueshen_stats":
            res = manager.stats()
            return [TextContent(type="text", text=json.dumps(res, ensure_ascii=False, indent=2))]

        return [
            TextContent(
                type="text",
                text=json.dumps({"success": False, "message": f"unknown tool: {name}"}, ensure_ascii=False),
            )
        ]
    except Exception as exc:
        return [
            TextContent(
                type="text",
                text=json.dumps({"success": False, "message": f"exception: {exc}"}, ensure_ascii=False),
            )
        ]


async def main():
    auto = AutoIngestor(manager)
    auto.start()
    async with mcp.server.stdio.stdio_server() as (read_stream, write_stream):
        init_opts = server.create_initialization_options()
        await server.run(read_stream, write_stream, init_opts)


if __name__ == "__main__":
    import asyncio

    asyncio.run(main())
